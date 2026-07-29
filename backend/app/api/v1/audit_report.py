# -*- coding: utf-8 -*-
"""
能源审计报告自动生成路由（符合 GB/T 17125《能源审计技术通则》）
- GET  /api/audit/report     ：生成能源审计报告（返回 JSON 结构化报告）
- GET  /api/audit/buildings  ：可审计建筑列表
- POST /api/audit/export     ：导出报告为 Word 文档（返回二进制流）

报告结构：
1. 基本信息：建筑名称、类型、面积、审计期间
2. 能耗概况：总能耗、能耗强度(kWh/㎡)、同类型建筑对比
3. 分项能耗：照明、空调、动力、特殊设备占比
4. 能效指标：COP 均值、负载率、运行时长
5. 节能潜力分析：识别低效设备、运行优化建议
6. 改造建议：具体改造措施、预期节能率、投资估算
7. 结论与建议

数据来源：从 fact_energy_records 按 building_id 聚合计算
节能潜力 = (实际能耗强度 - 行业先进值) × 面积
审计期间：最近 30 天（用 MAX(DATE(monitor_time)) 往前推）
"""
import io
import math
import logging
import datetime
from typing import Optional
from urllib.parse import quote

import pandas as pd
from fastapi import APIRouter, Request, HTTPException, Query
from fastapi.responses import JSONResponse, StreamingResponse
from pydantic import BaseModel, Field

from app.core.database import get_conn, run_in_thread, DBUnavailableError
from app.core.response_cache import cache_response
from app.core.route_error import handle_route_error

logger = logging.getLogger(__name__)
router = APIRouter()

# ===== 行业能耗基准（kWh/㎡·年），用于节能潜力计算 =====
# 依据：GB 50189《公共建筑节能设计标准》及行业能耗限额
# 字段：limit(限额) / advanced(先进值) / reasonable(合理值)
BENCHMARK_STANDARDS = {
    "OFFICE":     {"limit": 70, "advanced": 50, "reasonable": 65, "name": "办公楼"},
    "TEACHING":   {"limit": 45, "advanced": 30, "reasonable": 40, "name": "教学楼"},
    "LIBRARY":    {"limit": 55, "advanced": 35, "reasonable": 50, "name": "图书馆"},
    "CANTEEN":    {"limit": 80, "advanced": 55, "reasonable": 70, "name": "食堂"},
    "DORMITORY":  {"limit": 40, "advanced": 25, "reasonable": 35, "name": "公寓"},
    "LABORATORY": {"limit": 90, "advanced": 60, "reasonable": 80, "name": "科研实验"},
    "PLAZA":      {"limit": 20, "advanced": 15, "reasonable": 18, "name": "广场"},
    "CONFERENCE": {"limit": 70, "advanced": 50, "reasonable": 65, "name": "会议中心"},
}
DEFAULT_STANDARD = {"limit": 50, "advanced": 35, "reasonable": 45, "name": "其他"}

# param_type → 分项能耗大类映射
# 照明 / 空调 / 动力 / 特殊设备 四大类
PARAM_TYPE_CATEGORY = {
    "HVAC": "空调",
    "PRECISION_AC": "空调",
    "REFRIGERATION": "空调",
    "WATER_HEATER": "空调",
    "LIGHTING": "照明",
    "SOCKET": "动力",
    "VENTILATION": "动力",
    "PUMP": "动力",
    "EV_CHARGER": "特殊设备",
}

# 分项能耗大类顺序
ENERGY_CATEGORY_ORDER = ["空调", "照明", "动力", "特殊设备"]


def _safe_float(v, ndigits=2):
    """安全转换为 float，处理 NaN/Infinity/None"""
    try:
        if v is None:
            return None
        f = float(v)
        if math.isnan(f) or math.isinf(f):
            return None
        return round(f, ndigits)
    except (TypeError, ValueError):
        return None


def _get_standard(building_type: Optional[str]) -> dict:
    """根据建筑类型获取基准标准"""
    if not building_type:
        return DEFAULT_STANDARD.copy()
    return BENCHMARK_STANDARDS.get(building_type, DEFAULT_STANDARD).copy()


def _calc_grade(intensity: float, std: dict) -> str:
    """根据能耗强度与基准评级：A 先进 / B 合理 / C 落后 / D 超标"""
    if intensity < std["advanced"]:
        return "A"
    if intensity < std["reasonable"]:
        return "B"
    if intensity < std["limit"]:
        return "C"
    return "D"


GRADE_DESC = {
    "A": {"name": "先进", "desc": "实际能耗低于先进值"},
    "B": {"name": "合理", "desc": "实际能耗在先进值与合理值之间"},
    "C": {"name": "落后", "desc": "实际能耗在合理值与限额之间"},
    "D": {"name": "超标", "desc": "实际能耗超过限额"},
}


# ===== 数据采集 =====
def _fetch_building_meta(building_id: str) -> Optional[dict]:
    """获取建筑元信息"""
    with get_conn() as conn:
        df = pd.read_sql(
            "SELECT * FROM dim_buildings WHERE building_id = ?",
            conn,
            params=[building_id],
        )
    if df.empty:
        return None
    return dict(df.iloc[0])


def _fetch_all_buildings() -> pd.DataFrame:
    """获取全部建筑清单"""
    with get_conn() as conn:
        df = pd.read_sql(
            "SELECT building_id, building_name, building_type, total_area FROM dim_buildings ORDER BY building_id",
            conn,
        )
    return df


def _fetch_audit_period(building_id: str) -> dict:
    """获取审计期间：以 MAX(DATE(monitor_time)) 为终点，往前推 30 天"""
    with get_conn() as conn:
        row = conn.execute(
            """
            SELECT MAX(DATE(monitor_time)) AS end_day,
                   datetime(MAX(DATE(monitor_time)), '-30 days') AS start_day
            FROM fact_energy_records
            WHERE building_id = ?
            """,
            [building_id],
        ).fetchone()
    if row is None or not row["end_day"]:
        # 回退到当前日期
        today = datetime.date.today()
        return {
            "start": (today - datetime.timedelta(days=30)).isoformat(),
            "end": today.isoformat(),
            "days": 30,
        }
    return {
        "start": str(row["start_day"]),
        "end": str(row["end_day"]),
        "days": 30,
    }


def _fetch_building_energy_summary(building_id: str, period: dict) -> dict:
    """获取审计期间能耗汇总"""
    with get_conn() as conn:
        df = pd.read_sql(
            """
            SELECT
                SUM(elec_consumption) AS total_kwh,
                COUNT(*) AS record_cnt,
                COUNT(DISTINCT DATE(monitor_time)) AS day_cnt,
                AVG(elec_consumption) AS avg_hourly_kwh,
                MAX(elec_consumption) AS peak_hourly_kwh,
                SUM(CASE WHEN run_status != 'NORMAL' THEN 1 ELSE 0 END) AS abnormal_cnt
            FROM fact_energy_records
            WHERE building_id = ?
              AND DATE(monitor_time) >= ?
              AND DATE(monitor_time) <= ?
            """,
            conn,
            params=[building_id, period["start"], period["end"]],
        )
    if df.empty:
        return {}
    return dict(df.iloc[0])


def _fetch_subitem_energy(building_id: str, period: dict) -> pd.DataFrame:
    """获取分项能耗（按 param_type 聚合）"""
    with get_conn() as conn:
        df = pd.read_sql(
            """
            SELECT param_type,
                   SUM(elec_consumption) AS kwh,
                   COUNT(DISTINCT device_id) AS device_cnt
            FROM fact_energy_records
            WHERE building_id = ?
              AND DATE(monitor_time) >= ?
              AND DATE(monitor_time) <= ?
            GROUP BY param_type
            ORDER BY kwh DESC
            """,
            conn,
            params=[building_id, period["start"], period["end"]],
        )
    return df


def _fetch_daily_trend(building_id: str, period: dict) -> list:
    """获取审计期间每日能耗趋势（前端 daily_trend 字段）"""
    with get_conn() as conn:
        rows = conn.execute(
            """
            SELECT DATE(monitor_time) AS day, SUM(elec_consumption) AS kwh
            FROM fact_energy_records
            WHERE building_id = ?
              AND DATE(monitor_time) >= ?
              AND DATE(monitor_time) <= ?
            GROUP BY DATE(monitor_time)
            ORDER BY day
            """,
            [building_id, period["start"], period["end"]],
        ).fetchall()
    return [{"day": r["day"], "kwh": round(r["kwh"] or 0, 2)} for r in rows]


def _fetch_efficiency_metrics(building_id: str, period: dict) -> dict:
    """获取能效指标：COP 均值、负载率、运行时长"""
    with get_conn() as conn:
        df = pd.read_sql(
            """
            SELECT
                AVG(CASE WHEN cop IS NOT NULL THEN cop END) AS avg_cop,
                MIN(CASE WHEN cop IS NOT NULL THEN cop END) AS min_cop,
                MAX(CASE WHEN cop IS NOT NULL THEN cop END) AS max_cop,
                AVG(CASE WHEN loading_rate IS NOT NULL THEN loading_rate END) AS avg_loading_rate,
                COUNT(DISTINCT device_id) AS device_cnt,
                COUNT(DISTINCT DATE(monitor_time)) AS run_days
            FROM fact_energy_records
            WHERE building_id = ?
              AND DATE(monitor_time) >= ?
              AND DATE(monitor_time) <= ?
            """,
            conn,
            params=[building_id, period["start"], period["end"]],
        )
    if df.empty:
        return {}
    return dict(df.iloc[0])


def _fetch_low_efficiency_devices(building_id: str, period: dict, top_n: int = 10) -> pd.DataFrame:
    """识别低效设备：COP 均值偏低或负载率偏低"""
    with get_conn() as conn:
        df = pd.read_sql(
            """
            SELECT
                r.device_id,
                r.device_name,
                r.param_type,
                AVG(r.cop) AS avg_cop,
                AVG(r.loading_rate) AS avg_loading_rate,
                SUM(r.elec_consumption) AS total_kwh,
                COUNT(*) AS record_cnt,
                d.nominal_cop,
                d.rated_power
            FROM fact_energy_records r
            LEFT JOIN dim_devices d ON d.device_id = r.device_id
            WHERE r.building_id = ?
              AND DATE(r.monitor_time) >= ?
              AND DATE(r.monitor_time) <= ?
              AND r.cop IS NOT NULL
            GROUP BY r.device_id, r.device_name, r.param_type, d.nominal_cop, d.rated_power
            HAVING avg_cop IS NOT NULL
            ORDER BY avg_cop ASC
            LIMIT ?
            """,
            conn,
            params=[building_id, period["start"], period["end"], top_n],
        )
    return df


def _fetch_peer_comparison(building_type: str, period: dict) -> pd.DataFrame:
    """同类型建筑横向对比"""
    with get_conn() as conn:
        df = pd.read_sql(
            """
            SELECT b.building_id, b.building_name, b.total_area,
                   SUM(r.elec_consumption) AS kwh_30d
            FROM dim_buildings b
            LEFT JOIN fact_energy_records r
              ON r.building_id = b.building_id
              AND DATE(r.monitor_time) >= ?
              AND DATE(r.monitor_time) <= ?
            WHERE b.building_type = ?
            GROUP BY b.building_id, b.building_name, b.total_area
            ORDER BY kwh_30d DESC
            """,
            conn,
            params=[period["start"], period["end"], building_type],
        )
    return df


# ===== 报告组装 =====
def _build_audit_report(building_id: str) -> dict:
    """组装完整审计报告（JSON 结构）"""
    # 建筑元信息
    meta = _fetch_building_meta(building_id)
    if meta is None:
        return None

    building_type = meta.get("building_type") or ""
    building_name = meta.get("building_name") or building_id
    total_area = _safe_float(meta.get("total_area"), 2) or 0
    std = _get_standard(building_type)

    # 审计期间
    period = _fetch_audit_period(building_id)

    # 能耗概况
    energy = _fetch_building_energy_summary(building_id, period)
    total_kwh = _safe_float(energy.get("total_kwh"), 2) or 0
    day_cnt = int(energy.get("day_cnt") or 0)
    # 年化能耗强度：(30 天累计 / 实际天数) × 365 / 面积
    annual_kwh = (total_kwh / max(1, day_cnt)) * 365 if total_kwh > 0 else 0
    intensity = annual_kwh / total_area if total_area > 0 else 0
    grade = _calc_grade(intensity, std)

    # 节能潜力 = (实际强度 - 先进值) × 面积
    saving_potential_kwh = max(0.0, (intensity - std["advanced"]) * total_area)
    saving_potential_pct = (saving_potential_kwh / annual_kwh * 100) if annual_kwh > 0 else 0

    # 分项能耗
    subitem_df = _fetch_subitem_energy(building_id, period)
    subitem_breakdown = []
    category_kwh: dict = {cat: 0.0 for cat in ENERGY_CATEGORY_ORDER}
    total_sub_kwh = 0.0
    for _, r in subitem_df.iterrows():
        kwh = _safe_float(r["kwh"], 2) or 0
        cat = PARAM_TYPE_CATEGORY.get(str(r["param_type"]), "特殊设备")
        if cat not in category_kwh:
            category_kwh[cat] = 0.0
        category_kwh[cat] += kwh
        total_sub_kwh += kwh
        subitem_breakdown.append({
            "param_type": str(r["param_type"]),
            "category": cat,
            "kwh": round(kwh, 2),
            "device_cnt": int(r["device_cnt"]),
            "pct": round(kwh / total_kwh * 100, 2) if total_kwh > 0 else 0,
        })

    category_breakdown = []
    for cat in ENERGY_CATEGORY_ORDER:
        kwh = category_kwh.get(cat, 0)
        category_breakdown.append({
            "category": cat,
            "kwh": round(kwh, 2),
            "pct": round(kwh / total_sub_kwh * 100, 2) if total_sub_kwh > 0 else 0,
        })

    # 能效指标
    eff = _fetch_efficiency_metrics(building_id, period)
    avg_cop = _safe_float(eff.get("avg_cop"), 3)
    avg_loading_rate = _safe_float(eff.get("avg_loading_rate"), 2)
    run_days = int(eff.get("run_days") or 0)
    run_hours = run_days * 24  # 估算运行时长

    # 同类型建筑对比
    peer_df = _fetch_peer_comparison(building_type, period)
    peers = []
    for _, r in peer_df.iterrows():
        area = _safe_float(r["total_area"], 2) or 0
        kwh = _safe_float(r["kwh_30d"], 2) or 0
        peer_annual = (kwh / max(1, day_cnt)) * 365 if kwh > 0 else 0
        peer_intensity = peer_annual / area if area > 0 else 0
        peers.append({
            "building_id": str(r["building_id"]),
            "building_name": str(r["building_name"]),
            "total_area": area,
            "annual_kwh": round(peer_annual, 2),
            "intensity": round(peer_intensity, 2),
            "is_current": str(r["building_id"]) == building_id,
        })
    peers.sort(key=lambda x: x["intensity"])
    for i, p in enumerate(peers):
        p["rank"] = i + 1
    current_rank = next((p["rank"] for p in peers if p["is_current"]), None)
    peer_avg_intensity = round(
        sum(p["intensity"] for p in peers) / max(1, len(peers)), 2
    ) if peers else None
    peer_best_intensity = min((p["intensity"] for p in peers), default=None) if peers else None

    # 低效设备识别
    low_eff_df = _fetch_low_efficiency_devices(building_id, period, top_n=10)
    low_efficiency_devices = []
    for _, r in low_eff_df.iterrows():
        nominal_cop = _safe_float(r["nominal_cop"], 2)
        avg_cop_dev = _safe_float(r["avg_cop"], 3)
        # COP 衰减率 = (nominal - actual) / nominal
        degradation = None
        if nominal_cop and nominal_cop > 0 and avg_cop_dev is not None:
            degradation = round((nominal_cop - avg_cop_dev) / nominal_cop * 100, 1)
        low_efficiency_devices.append({
            "device_id": str(r["device_id"]),
            "device_name": str(r["device_name"]),
            "param_type": str(r["param_type"]),
            "avg_cop": avg_cop_dev,
            "nominal_cop": nominal_cop,
            "cop_degradation_pct": degradation,
            "avg_loading_rate": _safe_float(r["avg_loading_rate"], 2),
            "total_kwh": _safe_float(r["total_kwh"], 2),
        })

    # 节能潜力分析建议
    optimization_suggestions = _build_optimization_suggestions(
        grade, intensity, std, low_efficiency_devices, category_breakdown
    )

    # 改造建议
    renovation_measures = _build_renovation_measures(
        low_efficiency_devices, category_breakdown, saving_potential_kwh
    )

    # 结论
    conclusion = _build_conclusion(
        building_name, grade, intensity, std, saving_potential_kwh,
        saving_potential_pct, avg_cop, current_rank, len(peers)
    )

    report = {
        "report_meta": {
            "building_id": building_id,
            "building_name": building_name,
            "generated_at": datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            "standard": "GB/T 17125《能源审计技术通则》",
            "audit_period": period,
        },
        # ===== 前端期望的顶级字段（兼容层）=====
        "audit_period": {
            "days": day_cnt,
            "start": period.get("start"),
            "end": period.get("end"),
        },
        "basic_info": {
            "building_id": building_id,
            "building_name": building_name,
            "building_type": building_type,
            "building_type_name": std["name"],
            "total_area": total_area,
            "location_zone": meta.get("location_zone") or "",
        },
        "energy_overview": {
            "total_kwh": round(total_kwh, 2),
            "annual_kwh": round(annual_kwh, 2),
            "intensity": round(intensity, 2),
            "energy_intensity_kwh_per_m2": round(intensity, 2),  # 前端期望字段
            "unit": "kWh/㎡·年",
            "grade": grade,
            "grade_name": GRADE_DESC[grade]["name"],
            "grade_desc": GRADE_DESC[grade]["desc"],
            "avg_hourly_kwh": _safe_float(energy.get("avg_hourly_kwh"), 2),
            "peak_hourly_kwh": _safe_float(energy.get("peak_hourly_kwh"), 2),
            "abnormal_count": int(energy.get("abnormal_cnt") or 0),
            "record_count": int(energy.get("record_cnt") or 0),
            "day_count": day_cnt,
            "standard": std,
            "saving_potential_kwh": round(saving_potential_kwh, 2),
            "saving_potential_pct": round(saving_potential_pct, 2),
            "comparison": {
                "rating_name": GRADE_DESC[grade]["name"],
                "rating": grade,
            },
        },
        "subitem_energy": {
            "by_param_type": subitem_breakdown,
            "by_category": category_breakdown,
        },
        "efficiency_metrics": {
            "avg_cop": avg_cop,
            "min_cop": _safe_float(eff.get("min_cop"), 3),
            "max_cop": _safe_float(eff.get("max_cop"), 3),
            "avg_loading_rate": avg_loading_rate,
            "device_count": int(eff.get("device_cnt") or 0),
            "run_days": run_days,
            "run_hours": run_hours,
        },
        "peer_comparison": {
            "total_peers": len(peers),
            "current_rank": current_rank,
            "peer_avg_intensity": peer_avg_intensity,
            "peer_best_intensity": peer_best_intensity,
            "peers": peers,
        },
        "saving_potential": {
            "low_efficiency_devices": low_efficiency_devices,
            "optimization_suggestions": optimization_suggestions,
        },
        "renovation_suggestions": renovation_measures,
        "conclusion": conclusion,
        # ===== 前端期望的顶级字段（兼容层）=====
        "daily_trend": _fetch_daily_trend(building_id, period),
        "type_breakdown": [
            {"device_type": s["param_type"], "kwh": s["kwh"]}
            for s in subitem_breakdown
        ],
        "suggestions": optimization_suggestions + [
            m.get("measure", "") for m in renovation_measures if m.get("measure")
        ],
    }
    return report


def _build_optimization_suggestions(grade, intensity, std, low_eff_devices, categories) -> list:
    """运行优化建议"""
    suggestions = []
    if grade in ("C", "D"):
        suggestions.append(
            f"建筑能耗强度 {intensity:.1f} kWh/㎡·年，超过{'限额' if grade == 'D' else '合理值'}"
            f"（{std['limit'] if grade == 'D' else std['reasonable']}），建议优先开展系统性能调优。"
        )
    # 低效设备建议
    for dev in low_eff_devices[:3]:
        if dev.get("cop_degradation_pct") and dev["cop_degradation_pct"] > 20:
            suggestions.append(
                f"{dev['device_name']} COP 衰退 {dev['cop_degradation_pct']}%，"
                f"当前 COP={dev['avg_cop']}（额定 {dev['nominal_cop']}），建议安排深度维保或更换。"
            )
        elif dev.get("avg_loading_rate") is not None and dev["avg_loading_rate"] < 30:
            suggestions.append(
                f"{dev['device_name']} 平均负载率仅 {dev['avg_loading_rate']}%，"
                f"存在大马拉小车现象，建议优化台数控制或加装变频。"
            )
    # 分项能耗建议
    for cat in categories:
        if cat["pct"] > 50 and cat["category"] == "空调":
            suggestions.append(
                f"空调系统能耗占比 {cat['pct']:.1f}%，为最大耗能项，"
                f"建议优化冷冻水温度设定、加强冷凝器清洗。"
            )
    if not suggestions:
        suggestions.append("当前运行状态良好，建议保持现有策略并持续监测。")
    return suggestions


def _build_renovation_measures(low_eff_devices, categories, saving_kwh) -> list:
    """改造建议（措施、预期节能率、投资估算）"""
    measures = []

    # 措施1：低效设备更换
    severe_devs = [d for d in low_eff_devices if (d.get("cop_degradation_pct") or 0) > 30]
    if severe_devs:
        invest = len(severe_devs) * 8.0  # 单台约 8 万元
        measures.append({
            "measure": "高衰退设备整机更换",
            "target_devices": [d["device_name"] for d in severe_devs[:3]],
            "expected_saving_rate": 15.0,
            "expected_saving_kwh": round(saving_kwh * 0.4, 2),
            "investment_wan_yuan": round(invest, 1),
            "payback_years": round(invest / max(0.1, saving_kwh * 0.4 * 0.8 / 10000), 1),
            "priority": "高",
        })

    # 措施2：变频改造
    low_loading_devs = [d for d in low_eff_devices if (d.get("avg_loading_rate") or 100) < 40]
    if low_loading_devs:
        invest = len(low_loading_devs) * 3.0
        measures.append({
            "measure": "水泵/风机变频改造",
            "target_devices": [d["device_name"] for d in low_loading_devs[:3]],
            "expected_saving_rate": 8.0,
            "expected_saving_kwh": round(saving_kwh * 0.25, 2),
            "investment_wan_yuan": round(invest, 1),
            "payback_years": round(invest / max(0.1, saving_kwh * 0.25 * 0.8 / 10000), 1),
            "priority": "中",
        })

    # 措施3：照明 LED 改造
    lighting_pct = next((c["pct"] for c in categories if c["category"] == "照明" and c["kwh"] > 0), 0)
    if lighting_pct > 10:
        measures.append({
            "measure": "照明系统 LED 智能化改造",
            "target_devices": ["全部照明灯具"],
            "expected_saving_rate": 5.0,
            "expected_saving_kwh": round(saving_kwh * 0.15, 2),
            "investment_wan_yuan": round(saving_kwh * 0.15 * 0.3 / 10000, 1),
            "payback_years": 2.5,
            "priority": "中",
        })

    # 措施4：智控系统
    if saving_kwh > 0:
        measures.append({
            "measure": "能源管理智控系统部署",
            "target_devices": ["全楼宇"],
            "expected_saving_rate": 5.0,
            "expected_saving_kwh": round(saving_kwh * 0.2, 2),
            "investment_wan_yuan": 15.0,
            "payback_years": round(15.0 / max(0.1, saving_kwh * 0.2 * 0.8 / 10000), 1),
            "priority": "低",
        })

    return measures


def _build_conclusion(building_name, grade, intensity, std, saving_kwh, saving_pct,
                      avg_cop, rank, peer_total) -> str:
    """结论与建议"""
    grade_name = GRADE_DESC[grade]["name"]
    grade_desc = GRADE_DESC[grade]["desc"]
    parts = [
        f"本次审计对 {building_name} 近 30 天运行数据进行了系统分析。",
        f"建筑年化能耗强度为 {intensity:.1f} kWh/㎡·年，能效等级评定为 {grade_name}级"
        f"（{grade_desc}）。",
    ]
    if avg_cop is not None:
        parts.append(f"空调系统平均 COP={avg_cop}。")
    if rank and peer_total:
        parts.append(f"在同类型建筑中能耗强度排名第 {rank}/{peer_total}。")
    if saving_kwh > 0:
        parts.append(
            f"经测算，该建筑年节能潜力约 {saving_kwh:.0f} kWh（{saving_pct:.1f}%），"
            f"建议按优先级实施改造措施。"
        )
    else:
        parts.append("建筑能效已达行业先进水平，建议保持现有运行策略。")
    return "".join(parts)


# ===== 请求模型 =====
class ExportRequest(BaseModel):
    """导出报告请求"""
    building_id: str = Field(..., description="建筑ID")


# ===== 路由 =====
@router.get("/api/audit/buildings")
@run_in_thread
def list_auditable_buildings():
    """可审计建筑列表"""
    try:
        df = _fetch_all_buildings()
        if df.empty:
            return {"status": "success", "data": [], "message": "无建筑数据"}
        buildings = []
        for _, r in df.iterrows():
            std = _get_standard(r.get("building_type"))
            buildings.append({
                "building_id": str(r["building_id"]),
                "building_name": str(r["building_name"]),
                "building_type": str(r.get("building_type") or ""),
                "building_type_name": std["name"],
                "total_area": _safe_float(r.get("total_area"), 2),
            })
        return {
            "status": "success",
            "data": buildings,
            "total": len(buildings),
        }
    except DBUnavailableError:
        raise
    except Exception as e:
        return handle_route_error(e, logger, "可审计建筑列表查询")


@router.get("/api/audit/report")
@cache_response(ttl=300)  # 审计报告计算量大，缓存 5 分钟
@run_in_thread
def generate_audit_report(
    building_id: str = Query(..., description="建筑ID"),
):
    """生成能源审计报告（JSON 结构化）"""
    try:
        report = _build_audit_report(building_id)
        if report is None:
            raise HTTPException(status_code=404, detail=f"建筑不存在: {building_id}")
        return {"status": "success", "data": report}
    except HTTPException:
        raise
    except DBUnavailableError:
        raise
    except Exception as e:
        return handle_route_error(e, logger, "审计报告生成")


@router.post("/api/audit/export")
@run_in_thread
def export_audit_report(payload: ExportRequest):
    """导出审计报告为 Word 文档（返回二进制流）"""
    try:
        report = _build_audit_report(payload.building_id)
        if report is None:
            raise HTTPException(status_code=404, detail=f"建筑不存在: {payload.building_id}")

        file_stream = _render_docx(report)
        filename = f"能源审计报告_{report['basic_info']['building_name']}_{datetime.datetime.now().strftime('%Y%m%d')}.docx"
        encoded_filename = quote(filename)
        return StreamingResponse(
            file_stream,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename*=UTF-8''{encoded_filename}"},
        )
    except HTTPException:
        raise
    except DBUnavailableError:
        raise
    except Exception as e:
        return handle_route_error(e, logger, "审计报告导出")


# ===== Word 渲染 =====
def _render_docx(report: dict) -> io.BytesIO:
    """将审计报告渲染为 Word 文件流"""
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn

    doc = Document()
    main_color = RGBColor(30, 58, 138)
    gray_color = RGBColor(100, 100, 100)

    def set_style(run, size=10.5, bold=False, color=None):
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        run.font.size = Pt(size)
        run.bold = bold
        if color:
            run.font.color.rgb = color

    def add_heading(text, level=1):
        h = doc.add_heading("", level=level)
        run = h.add_run(text)
        set_style(run, size=16 if level == 1 else 14, bold=True, color=main_color)
        return h

    def add_para(text, size=10.5, bold=False, color=None, align=None):
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.5
        if align:
            p.alignment = align
        run = p.add_run(text)
        set_style(run, size=size, bold=bold, color=color)
        return p

    # ===== 标题 =====
    title = doc.add_heading("", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("建筑能源审计报告")
    set_style(run, size=22, bold=True, color=main_color)

    meta = report["report_meta"]
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = sub.add_run(
        f"报告依据：{meta['standard']}　|　生成时间：{meta['generated_at']}"
    )
    set_style(run_sub, size=9, color=gray_color)

    # ===== 一、基本信息 =====
    add_heading("一、基本信息", level=1)
    basic = report["basic_info"]
    period = meta["audit_period"]
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    for label, val in [
        ("建筑名称", basic["building_name"]),
        ("建筑编号", basic["building_id"]),
        ("建筑类型", f"{basic['building_type_name']}（{basic['building_type']}）"),
        ("建筑面积", f"{basic['total_area']} ㎡"),
        ("所在区域", basic.get("location_zone") or "—"),
        ("审计期间", f"{period['start']} 至 {period['end']}（近 {period['days']} 天）"),
    ]:
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = str(val)

    # ===== 二、能耗概况 =====
    add_heading("二、能耗概况", level=1)
    ov = report["energy_overview"]
    add_para(
        f"审计期间总能耗 {ov['total_kwh']} kWh，年化能耗约 {ov['annual_kwh']} kWh，"
        f"能耗强度 {ov['intensity']} {ov['unit']}，"
        f"能效等级：{ov['grade_name']}级（{ov['grade_desc']}）。",
        size=11,
    )
    add_para(
        f"行业基准——限额 {ov['standard']['limit']} / 合理值 {ov['standard']['reasonable']} / "
        f"先进值 {ov['standard']['advanced']}（单位：{ov['unit']}）。",
        size=10, color=gray_color,
    )
    add_para(
        f"理论节能潜力 {ov['saving_potential_kwh']} kWh/年（占比 {ov['saving_potential_pct']}%）。",
        size=11, bold=True,
    )

    # 同类建筑对比
    peer = report["peer_comparison"]
    if peer["total_peers"] > 0:
        add_para(
            f"同类型建筑共 {peer['total_peers']} 栋，本建筑能耗强度排名第 "
            f"{peer['current_rank']}，同类均值 {peer['peer_avg_intensity']} {ov['unit']}，"
            f"同类最优 {peer['peer_best_intensity']} {ov['unit']}。",
            size=11,
        )

    # ===== 三、分项能耗 =====
    add_heading("三、分项能耗", level=1)
    add_para("按用能大类占比：", size=11)
    t3 = doc.add_table(rows=1, cols=3)
    t3.style = "Table Grid"
    t3.rows[0].cells[0].text = "用能大类"
    t3.rows[0].cells[1].text = "能耗(kWh)"
    t3.rows[0].cells[2].text = "占比"
    for cat in report["subitem_energy"]["by_category"]:
        cells = t3.add_row().cells
        cells[0].text = cat["category"]
        cells[1].text = str(cat["kwh"])
        cells[2].text = f"{cat['pct']}%"

    add_para("按设备类型明细：", size=11)
    t4 = doc.add_table(rows=1, cols=4)
    t4.style = "Table Grid"
    for i, h in enumerate(["设备类型", "用能大类", "能耗(kWh)", "占比"]):
        t4.rows[0].cells[i].text = h
    for item in report["subitem_energy"]["by_param_type"]:
        cells = t4.add_row().cells
        cells[0].text = item["param_type"]
        cells[1].text = item["category"]
        cells[2].text = str(item["kwh"])
        cells[3].text = f"{item['pct']}%"

    # ===== 四、能效指标 =====
    add_heading("四、能效指标", level=1)
    eff = report["efficiency_metrics"]
    add_para(
        f"空调系统 COP 均值 {eff['avg_cop']}（区间 {eff['min_cop']} ~ {eff['max_cop']}），"
        f"平均负载率 {eff['avg_loading_rate']}%，"
        f"运行天数 {eff['run_days']} 天（约 {eff['run_hours']} 小时）。",
        size=11,
    )

    # ===== 五、节能潜力分析 =====
    add_heading("五、节能潜力分析", level=1)
    low_devs = report["saving_potential"]["low_efficiency_devices"]
    if low_devs:
        add_para(f"识别出 {len(low_devs)} 台低效设备（按 COP 升序）：", size=11)
        t5 = doc.add_table(rows=1, cols=5)
        t5.style = "Table Grid"
        for i, h in enumerate(["设备名称", "当前COP", "额定COP", "衰退率", "负载率"]):
            t5.rows[0].cells[i].text = h
        for d in low_devs:
            cells = t5.add_row().cells
            cells[0].text = d["device_name"]
            cells[1].text = str(d["avg_cop"])
            cells[2].text = str(d["nominal_cop"])
            cells[3].text = f"{d.get('cop_degradation_pct') or '—'}%"
            cells[4].text = f"{d.get('avg_loading_rate') or '—'}%"
    else:
        add_para("未发现明显低效设备。", size=11)

    add_para("运行优化建议：", size=11, bold=True)
    for s in report["saving_potential"]["optimization_suggestions"]:
        add_para(f"• {s}", size=10.5)

    # ===== 六、改造建议 =====
    add_heading("六、改造建议", level=1)
    measures = report["renovation_suggestions"]
    if measures:
        t6 = doc.add_table(rows=1, cols=6)
        t6.style = "Table Grid"
        for i, h in enumerate(["改造措施", "预期节能率", "节能kWh", "投资(万元)", "回收期(年)", "优先级"]):
            t6.rows[0].cells[i].text = h
        for m in measures:
            cells = t6.add_row().cells
            cells[0].text = m["measure"]
            cells[1].text = f"{m['expected_saving_rate']}%"
            cells[2].text = str(m["expected_saving_kwh"])
            cells[3].text = str(m["investment_wan_yuan"])
            cells[4].text = str(m["payback_years"])
            cells[5].text = m["priority"]
    else:
        add_para("暂无改造建议。", size=11)

    # ===== 七、结论与建议 =====
    add_heading("七、结论与建议", level=1)
    add_para(report["conclusion"], size=11)

    # ===== 落款 =====
    sign = doc.add_paragraph()
    sign.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    sign.add_run("\n\n特此报告，敬请审阅。\n\n")
    run_team = sign.add_run("能源审计工作组")
    set_style(run_team, size=12, bold=True)
    run_date = sign.add_run(f"\n{datetime.datetime.now().strftime('%Y年%m月%d日')}")
    set_style(run_date, size=11)

    # 页脚
    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_footer = footer.add_run(
        f"--- 本报告依据 {meta['standard']} 自动生成 | 内部文件 严禁外泄 ---"
    )
    set_style(run_footer, size=9, color=RGBColor(150, 150, 150))

    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream
