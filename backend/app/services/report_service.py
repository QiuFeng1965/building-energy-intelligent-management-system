# -*- coding: utf-8 -*-
"""
报告服务层
- generate_weekly_report_content：生成纯文本周报内容（供 AI 工具调用，返回给 LLM）
- generate_weekly_report_docx：生成 Word 文件流（供 /api/report/weekly_ai 路由下载）
两者共享数据采集与 AI 诊断逻辑，避免重复实现。
"""
import io
import datetime
import logging

import pandas as pd
from openai import AsyncOpenAI

from app.core.config import AI_API_KEY, AI_BASE_URL, MODEL_TEXT
from app.core.database import get_conn

logger = logging.getLogger(__name__)


def _fetch_weekly_kpi(building_type: str = None):
    """拉取近 7 天核心 KPI（total_elec / abnormal_count / avg_cop / avg_power）"""
    max_date = datetime.datetime.now().strftime('%Y-%m-%d')
    with get_conn() as conn:
        kpi_df = pd.read_sql(
            "SELECT SUM(elec_consumption) as total_elec, "
            "SUM(CASE WHEN run_status != 'NORMAL' THEN 1 ELSE 0 END) as alarms "
            "FROM fact_energy_records "
            "WHERE DATE(monitor_time) <= ? AND DATE(monitor_time) >= date(?, '-7 days')",
            conn, params=[max_date, max_date]
        )
        cop_df = pd.read_sql(
            "SELECT AVG(cop) as cop FROM fact_energy_records "
            "WHERE DATE(monitor_time) <= ? AND DATE(monitor_time) >= date(?, '-7 days') "
            "AND param_type = 'HVAC'",
            conn, params=[max_date, max_date]
        )
    total_elec = int(kpi_df['total_elec'].iloc[0] or 0)
    abnormal_count = int(kpi_df['alarms'].iloc[0] or 0)
    avg_cop = round(cop_df['cop'].iloc[0] or 4.5, 2)
    avg_power = round(total_elec / (7 * 24), 1) if total_elec > 0 else 0
    return {
        "total_elec": total_elec,
        "abnormal_count": abnormal_count,
        "avg_cop": avg_cop,
        "avg_power": avg_power,
        "max_date": max_date,
    }


async def _generate_ai_diagnosis(kpi: dict, issue: str = None) -> str:
    """调大模型生成诊断文本"""
    prompt = f"""
    你现在是【擎翼数字中枢】的 AI 首席能效工程师。请基于以下底层运行数据，直接输出【深度诊断结论】与【维保工单建议】。
    【本周实时数据】：
    - 耗电总量: {kpi['total_elec']:,} kWh | 负荷功率: {kpi['avg_power']} kW | 告警频次: {kpi['abnormal_count']} 次 | 系统 COP: {kpi['avg_cop']}
    {f"- 重点关注的异常问题：{issue}" if issue else ""}

    【输出要求】：
    1. 拒绝客套话。严禁出现"尊敬的领导"等字眼。
    2. 【深度诊断】：分析系统运行效率（如：COP 是否达标）。
    3. 【维保工单】：输出 2-3 条具备可操作性的工单建议（包含 [任务名称]、[触发原因]、[执行动作]）。
    4. 语言风格：专业、精炼、硬核。
    """
    try:
        client = AsyncOpenAI(api_key=AI_API_KEY, base_url=AI_BASE_URL)
        response = await client.chat.completions.create(
            model=MODEL_TEXT,
            messages=[{"role": "user", "content": prompt}],
            temperature=0.3,
            max_tokens=2000
        )
        return response.choices[0].message.content
    except Exception as e:
        logger.exception(f"AI 诊断生成失败: {e}")
        return (
            "【诊断引擎暂离，启用规则引擎兜底】\n"
            "诊断：COP 偏低，系统存在非计划性停机风险。\n"
            "工单 1：[冷凝器清洗] 因换热效率下降，建议检查水侧结垢情况。"
        )


async def generate_weekly_report_content(building_type: str = None, issue: str = None) -> str:
    """生成纯文本周报内容（供 AI 工具调用）"""
    try:
        kpi = _fetch_weekly_kpi(building_type)
        ai_analysis = await _generate_ai_diagnosis(kpi, issue)
        content = (
            f"=== 建筑能源智能管理与运营优化周报 ===\n"
            f"报告生成时间：{datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n"
            f"建筑类型：{building_type or '全校区'}\n\n"
            f"【一、本周底层真实数据萃取】\n"
            f"- 7天累计总耗电量：{kpi['total_elec']:,} kWh\n"
            f"- 7天累计异常次数：{kpi['abnormal_count']} 次\n"
            f"- 平台周平均功率：{kpi['avg_power']} kW\n"
            f"- 机组能效比(COP)：{kpi['avg_cop']}\n\n"
            f"【二、AI 深度能效诊断与维保工单建议】\n{ai_analysis}\n\n"
            f"【三、落实\"双碳\"战略建议】\n"
            f"1. 加强异常设备排查，重点检查告警机组；\n"
            f"2. 优化空调系统运行策略，提升核心机组能效比；\n"
            f"3. 引入智能化能耗管理系统，实现精细化管理。\n"
        )
        return content
    except Exception as e:
        logger.exception(f"生成周报内容失败: {e}")
        return "周报生成失败，请稍后重试或联系管理员"


async def generate_weekly_report_docx() -> io.BytesIO:
    """生成 Word 文件流（供 /api/report/weekly_ai 路由下载）"""
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn

    kpi = _fetch_weekly_kpi()
    ai_analysis = await _generate_ai_diagnosis(kpi)

    doc = Document()
    main_color = RGBColor(30, 58, 138)

    def set_style(run, size=10.5, bold=False, color=None):
        run.font.name = 'Times New Roman'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        run.font.size = Pt(size)
        run.bold = bold
        if color:
            run.font.color.rgb = color

    # 标题
    title = doc.add_heading('', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('建筑能源智能管理与运营优化周报')
    set_style(run, size=22, bold=True, color=main_color)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_meta = meta.add_run(f"报告生成时间：{datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')} | 数据源：Enterprise_DB")
    set_style(run_meta, size=9, color=RGBColor(100, 100, 100))

    # 数据表格
    h1 = doc.add_heading('', level=1)
    run_h1 = h1.add_run('一、 本周底层真实数据萃取')
    set_style(run_h1, size=18, bold=True, color=main_color)

    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    table.rows[0].cells[0].text = '监控维度'
    table.rows[0].cells[1].text = '本周运行数值'
    for label, val in [('7天累计总耗电量', f"{kpi['total_elec']:,} kWh"),
                       ('7天累计异常次数', f"{kpi['abnormal_count']} 次"),
                       ('平台周平均功率', f"{kpi['avg_power']} kW"),
                       ('机组能效比(COP)', f"{kpi['avg_cop']}")]:
        row = table.add_row().cells
        row[0].text, row[1].text = label, val

    # AI 诊断
    h2 = doc.add_heading('', level=1)
    run_h2 = h2.add_run('二、 AI 深度能效诊断与维保工单建议')
    set_style(run_h2, size=18, bold=True, color=main_color)
    p_ai = doc.add_paragraph()
    p_ai.paragraph_format.line_spacing = 1.5
    run_ai = p_ai.add_run(ai_analysis)
    set_style(run_ai, size=11)

    # 双碳建议
    h3 = doc.add_heading('', level=1)
    run_h3 = h3.add_run('三、 落实"双碳"战略建议与工作总结')
    set_style(run_h3, size=18, bold=True, color=main_color)

    p_summary = doc.add_paragraph()
    p_summary.paragraph_format.line_spacing = 1.5
    summary_text = (
        "结合国家\"双碳\"战略目标，针对上述情况，提出以下建议：\n"
        "一是加强对异常设备的排查与维护；\n"
        "二是优化空调系统运行策略，提升核心机组的能效比；\n"
        "三是逐步引入智能化能耗管理系统，实现精细化管理。"
    )
    run_sum = p_summary.add_run(summary_text)
    set_style(run_sum, size=11)

    # 落款
    sign = doc.add_paragraph()
    sign.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    sign.add_run("\n\n特此报告，敬请审阅。\n\n")
    run_team = sign.add_run("中建八局能效管理团队")
    set_style(run_team, size=12, bold=True)
    run_date = sign.add_run(f"\n{datetime.datetime.now().strftime('%Y年%m月%d日')}")
    set_style(run_date, size=11)

    # 页脚
    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_footer = footer.add_run("--- 本报告由 擎翼数字中枢 AI Agent 引擎自动生成 | 内部文件 严禁外泄 ---")
    set_style(run_footer, size=9, color=RGBColor(150, 150, 150))

    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream
